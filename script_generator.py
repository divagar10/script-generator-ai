import google.generativeai as genai
import datetime
from docx import Document
genai.configure(api_key="AIzaSyDcc1s1OF-U7Dr6n4s5MVzHJbzyzfZxVrY")
model = genai.GenerativeModel("gemini-1.5-flash")
video_type = input("Choose the type of content (YouTube video / Instagram reel / YouTube Shorts): ").strip().lower()
topic = input("Enter your video topic: ")
if "short" in video_type or "reel" in video_type:
    prompt = f"""
    Generate an engaging script for a 1-minute {video_type} about "{topic}".
    The script should have an attention-grabbing hook in the first few seconds,
    use casual and high-energy tone, include visual/action suggestions,
    and end with a strong call to action.
    """
else:
    prompt = f"""
    Generate a detailed script for a YouTube video about "{topic}".
    The script should last about 10 minutes and include:
    - Hook in the intro
    - Structured explanation in sections
    - Engaging storytelling tone
    - Visual suggestions for each section
    - Call to action at the end
    Format the response as a script with timestamps.
    """
response = model.generate_content(prompt)
script_text = response.text.strip()
print("\n--- Generated Script ---\n")
print(script_text)
timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
txt_path = f"youtube_script_{timestamp}.txt"
with open(txt_path, "w", encoding="utf-8") as f:
    f.write(script_text)
doc = Document()
doc.add_heading("Generated Script", 0)
doc.add_paragraph(script_text)
docx_path = f"youtube_script_{timestamp}.docx"
doc.save(docx_path)

print(f"\n✅ Script exported as:\n📄 {txt_path}\n📄 {docx_path}")
